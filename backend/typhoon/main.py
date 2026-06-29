import os
import io
import base64
import logging
import torch
from contextlib import asynccontextmanager
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import Optional, List
from PIL import Image
from transformers import (
    AutoTokenizer,
    AutoModelForCausalLM,
    AutoModelForImageTextToText,
    AutoProcessor,
    BitsAndBytesConfig,
    TextIteratorStreamer,
)
import threading
from fastapi.responses import StreamingResponse

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─── Model IDs ────────────────────────────────────────────────────────────────
CHAT_MODEL_ID = "typhoon-ai/typhoon2.5-qwen3-4b"
OCR_MODEL_ID  = "typhoon-ai/typhoon-ocr1.5-2b"

# ─── Global holders ───────────────────────────────────────────────────────────
models = {}

# Create cache directory if it doesn't exist
CACHE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".cache")
os.makedirs(CACHE_DIR, exist_ok=True)

# ─── OCR Prompt ───────────────────────────────────────────────────────────────
OCR_PROMPT = """Extract all text from the image.
Instructions:
- Only return the clean Markdown.
- Do not include any explanation or extra text.
- You must include all information on the page.
Formatting Rules:
- Tables: Render tables using <table>...</table> in clean HTML format.
- Equations: Render equations using LaTeX syntax with inline ($...$) and block ($$...$$).
- Images/Charts/Diagrams: Wrap any clearly defined visual areas in:
  <figure> Describe in Thai. </figure>
- Page Numbers: Wrap page numbers in <page_number>...</page_number>.
- Checkboxes: Use ☐ for unchecked and ☑ for checked boxes."""


def resize_if_needed(img: Image.Image, max_size: int = 1024) -> Image.Image:
    width, height = img.size
    if max(width, height) > max_size:
        scale = max_size / float(max(width, height))
        new_size = (int(width * scale), int(height * scale))
        img = img.resize(new_size, Image.Resampling.LANCZOS)
    return img


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Load models on startup with CPU fallback for compatibility."""
    
    # 1. OCR Model
    logger.info("Loading Typhoon OCR model...")
    try:
        # Use 4-bit quantization for OCR to save VRAM
        quant_config_ocr = BitsAndBytesConfig(
            load_in_4bit=True,
            bnb_4bit_compute_dtype=torch.float16,
            bnb_4bit_quant_type="nf4",
        )
        models["ocr_model"] = AutoModelForImageTextToText.from_pretrained(
            OCR_MODEL_ID,
            quantization_config=quant_config_ocr,
            device_map="auto",
        )
        models["ocr_processor"] = AutoProcessor.from_pretrained(OCR_MODEL_ID)
        logger.info("✅ OCR model loaded on GPU (4-bit).")
    except Exception as e:
        logger.warning(f"⚠️ GPU failed, falling back to CPU: {e}")
        try:
            models["ocr_model"] = AutoModelForImageTextToText.from_pretrained(
                OCR_MODEL_ID,
                torch_dtype=torch.float32,
                device_map={"": "cpu"},
            )
            models["ocr_processor"] = AutoProcessor.from_pretrained(OCR_MODEL_ID)
            logger.info("✅ OCR model loaded on CPU.")
        except Exception as e2:
            logger.error(f"❌ Failed to load OCR even on CPU: {e2}")

    # 2. Chat Model
    logger.info("Loading Typhoon 2.5 chat model...")
    try:
        # Correct way to load 4-bit using BitsAndBytesConfig
        quant_config = BitsAndBytesConfig(
            load_in_4bit=True,
            bnb_4bit_compute_dtype=torch.bfloat16,
            bnb_4bit_quant_type="nf4",
            bnb_4bit_use_double_quant=True,
        )
        
        models["chat_tokenizer"] = AutoTokenizer.from_pretrained(CHAT_MODEL_ID)
        models["chat_model"] = AutoModelForCausalLM.from_pretrained(
            CHAT_MODEL_ID,
            quantization_config=quant_config,
            device_map="auto",
        )
        logger.info("✅ Chat model loaded on GPU (4-bit).")
    except Exception as e:
        logger.warning(f"⚠️ GPU/4-bit failed, falling back to CPU: {e}")
        try:
            models["chat_model"] = AutoModelForCausalLM.from_pretrained(
                CHAT_MODEL_ID,
                torch_dtype=torch.float32,
                device_map={"": "cpu"},
            )
            logger.info("✅ Chat model loaded on CPU.")
        except Exception as e2:
            logger.error(f"❌ Failed to load Chat even on CPU: {e2}")

    yield
    models.clear()


app = FastAPI(title="Typhoon AI Platform", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─── Schemas ──────────────────────────────────────────────────────────────────
class ChatMessage(BaseModel):
    role: str   # "user" | "assistant" | "system"
    content: str


class ChatRequest(BaseModel):
    messages: List[ChatMessage]
    system_prompt: Optional[str] = None
    max_new_tokens: int = 1024
    temperature: float = 0.6
    top_p: float = 0.95


class ResumeAnalysisRequest(BaseModel):
    ocr_text: str
    job_description: Optional[str] = None


# ─── Routes ───────────────────────────────────────────────────────────────────
@app.get("/health")
def health():
    return {
        "status": "ok",
        "ocr_model": "ocr_model" in models,
        "chat_model": "chat_model" in models,
    }



# ─── Shared OCR helper (used by /ocr endpoint and /api/analyze) ───────────────
def _run_ocr_on_image(image: Image.Image, label: str = "") -> str:
    """
    Run Typhoon OCR (typhoon-ocr1.5-2b) on a PIL Image.
    The processor expects the image as a data URI (base64) inside the prompt,
    NOT as a raw PIL Image argument.
    """
    import base64, io as _bio

    proc  = models["ocr_processor"]
    model = models["ocr_model"]

    # Convert PIL Image to base64 JPEG data URI
    buf = _bio.BytesIO()
    image.save(buf, format="JPEG", quality=95)
    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    data_uri = f"data:image/jpeg;base64,{b64}"

    # Build the message in the format the processor expects
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {"url": data_uri}},
                {"type": "text",      "text": OCR_PROMPT},
            ],
        }
    ]

    prompt_text = proc.apply_chat_template(
        messages, tokenize=False, add_generation_prompt=True
    )
    inputs = proc(text=prompt_text, images=image, return_tensors="pt")
    # Move tensors to the same device as the model
    device = next(model.parameters()).device
    inputs = {k: v.to(device) for k, v in inputs.items()}

    with torch.no_grad():
        output_ids = model.generate(
            **inputs,
            max_new_tokens=1024,
            do_sample=False,
        )

    # Decode only newly generated tokens
    input_len = inputs["input_ids"].shape[1]
    new_ids   = output_ids[:, input_len:]
    text = proc.batch_decode(new_ids, skip_special_tokens=True)[0].strip()

    logger.info(f"[OCR] {label}: {len(text)} chars | {text[:80]!r}")
    return text


@app.post("/ocr")

async def ocr_endpoint(file: UploadFile = File(...)):
    """Extract text from images (OCR) or Documents (PDF/Docx)."""
    filename = file.filename.lower()
    content = await file.read()
    
    # 1. Handle PDF
    if filename.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return {"text": text, "type": "pdf"}
        except Exception as e:
            raise HTTPException(500, f"PDF Error: {e}")

    # 2. Handle DOCX
    elif filename.endswith(".docx"):
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            text = "\n".join([para.text for para in doc.paragraphs])
            return {"text": text, "type": "docx"}
        except Exception as e:
            raise HTTPException(500, f"Docx Error: {e}")

    # 3. Handle Images (Typhoon OCR)
    elif filename.endswith((".png", ".jpg", ".jpeg")):
        if "ocr_model" not in models:
            return {"text": f"OCR Simulated for {filename}", "type": "image"}
            
        try:
            from PIL import Image
            import io
            image = Image.open(io.BytesIO(content)).convert("RGB")
            image = resize_if_needed(image)
            text = _run_ocr_on_image(image, filename)
            return {"text": text, "markdown": text, "type": "image"}
        except Exception as e:
            logger.exception("OCR failed")
            raise HTTPException(500, str(e))

    
    else:
        try:
            return {"text": content.decode("utf-8"), "type": "text"}
        except:
            raise HTTPException(400, "Unsupported file type")


@app.post("/chat")
async def chat_endpoint(req: ChatRequest):
    """General chatbot via Typhoon 2.5 with Streaming support."""
    if "chat_model" not in models:
        raise HTTPException(503, "Chat model not loaded")

    try:
        tokenizer = models["chat_tokenizer"]
        model     = models["chat_model"]

        system_content = (
            req.system_prompt
            or (
                "You are a helpful AI assistant named Typhoon created by SCB 10X. "
                "You respond in the language the user uses (Thai or English). "
                "Be concise, friendly, and accurate."
            )
        )

        messages = [{"role": "system", "content": system_content}]
        for m in req.messages:
            messages.append({"role": m.role, "content": m.content})

        inputs = tokenizer.apply_chat_template(
            messages,
            add_generation_prompt=True,
            return_tensors="pt",
            return_dict=True,
        ).to(model.device)

        streamer = TextIteratorStreamer(tokenizer, skip_prompt=True, skip_special_tokens=True)
        
        generation_kwargs = dict(
            **inputs,
            streamer=streamer,
            max_new_tokens=req.max_new_tokens,
            do_sample=True,
            temperature=req.temperature,
            top_p=req.top_p,
            repetition_penalty=1.05,
        )

        # Run generation in a separate thread
        thread = threading.Thread(target=model.generate, kwargs=generation_kwargs)
        thread.start()

        def generate_and_stream():
            for new_text in streamer:
                yield new_text

        return StreamingResponse(generate_and_stream(), media_type="text/plain")

    except Exception as e:
        logger.exception("Chat failed")
        raise HTTPException(500, f"Chat failed: {str(e)}")

# --- Knowledge Management ---

@app.get("/list-docs")
async def list_docs():
    """List text files in the sample docs directory."""
    base_dir = r"c:\Users\Win11\Desktop\La\create_sample_docs\company_docs"
    try:
        if not os.path.exists(base_dir):
            return {"files": []}
        files = [f for f in os.listdir(base_dir) if f.endswith(('.txt', '.md'))]
        return {"files": files}
    except Exception as e:
        raise HTTPException(500, str(e))

@app.get("/list-resumes")
async def list_resumes():
    """List image files in the resumes folder."""
    base_dir = r"c:\Users\Win11\Desktop\Typhoon\backend\resumes"
    try:
        import os
        if not os.path.exists(base_dir):
            return {"files": []}
        files = [f for f in os.listdir(base_dir) if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
        return {"files": files}
    except Exception as e:
        raise HTTPException(500, str(e))

@app.get("/get-resume/{filename}")
async def get_resume(filename: str):
    """Get the binary content of a resume file."""
    base_dir = r"c:\Users\Win11\Desktop\Typhoon\backend\resumes"
    file_path = os.path.join(base_dir, filename)
    if not os.path.exists(file_path):
        raise HTTPException(404, "File not found")
    from fastapi.responses import FileResponse
    return FileResponse(file_path)

@app.get("/list-jobs")
async def list_jobs():
    """List directory names in the jobs folder."""
    base_dir = r"c:\Users\Win11\Desktop\Typhoon\backend\jobs"
    try:
        import os
        if not os.path.exists(base_dir):
            return {"jobs": []}
        # Only list directories (which represent job positions)
        jobs = [d for d in os.listdir(base_dir) if os.path.isdir(os.path.join(base_dir, d))]
        return {"jobs": jobs}
    except Exception as e:
        raise HTTPException(500, str(e))

@app.get("/get-doc-path")
async def get_doc_path(path: str):
    """Read content of a file by relative path inside backend/ directory."""
    base_dir = r"c:\Users\Win11\Desktop\Typhoon\backend"
    target_path = os.path.join(base_dir, path)

    # Path traversal protection
    if not os.path.abspath(target_path).startswith(os.path.abspath(base_dir)):
        raise HTTPException(403, "Access denied")

    if not os.path.exists(target_path):
        raise HTTPException(404, f"File not found: {path}")

    try:
        with open(target_path, "r", encoding="utf-8") as f:
            return {"content": f.read()}
    except Exception as e:
        raise HTTPException(500, str(e))


@app.get("/get-doc/{filename}")
async def get_doc(filename: str):
    """Read content of a specific knowledge file."""
    base_dir = r"c:\Users\Win11\Desktop\La\create_sample_docs\company_docs"
    target_path = os.path.join(base_dir, filename)
    
    # Path traversal protection
    if not os.path.abspath(target_path).startswith(os.path.abspath(base_dir)):
        raise HTTPException(403, "Access denied")
        
    try:
        with open(target_path, "r", encoding="utf-8") as f:
            return {"content": f.read()}
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/analyze-resume")
async def analyze_resume(req: ResumeAnalysisRequest):
    """Analyze resume text (from OCR) using Typhoon 2.5."""
    if "chat_model" not in models:
        raise HTTPException(503, "Chat model not loaded")

    system_prompt = (
        "You are an expert HR recruiter and resume analyst. "
        "Analyze the provided resume text and give structured, insightful feedback in Thai. "
        "Cover: 1) ข้อมูลส่วนตัวและสรุปโปรไฟล์ 2) ทักษะและความสามารถ 3) ประสบการณ์การทำงาน "
        "4) การศึกษา 5) จุดแข็งของผู้สมัคร 6) จุดที่ควรพัฒนา 7) คะแนนความเหมาะสม (1-10) "
        "และ 8) ข้อเสนอแนะ. Format your response clearly with headers."
    )

    if req.job_description:
        user_content = (
            f"กรุณาวิเคราะห์ resume นี้ให้ละเอียด:\n\n{req.ocr_text}\n\n"
            f"โดยเทียบกับ Job Description นี้:\n{req.job_description}"
        )
    else:
        user_content = f"กรุณาวิเคราะห์ resume นี้ให้ละเอียด:\n\n{req.ocr_text}"

    chat_req = ChatRequest(
        messages=[ChatMessage(role="user", content=user_content)],
        system_prompt=system_prompt,
        max_new_tokens=2048,
    )
    return await chat_endpoint(chat_req)


# ─── Llama-style Dashboard API ────────────────────────────────────────────────

JOBS_BASE   = r"c:\Users\Win11\Desktop\Typhoon\backend\jobs"
RESUME_BASE = r"c:\Users\Win11\Desktop\Typhoon\backend\resumes"

@app.get("/api/roles")
async def api_roles():
    """List all job roles (same as Llama /api/roles)."""
    if not os.path.exists(JOBS_BASE):
        return []
    return [d for d in os.listdir(JOBS_BASE) if os.path.isdir(os.path.join(JOBS_BASE, d))]


import re as _re
import json as _json

def _parse_criteria(text: str) -> dict:
    criteria_map = {}
    for line in text.split("\n"):
        m = _re.search(r"(\d+\..+?)\s*\((\d+)\s*คะแนน\)", line)
        if m:
            key = f"cat_{len(criteria_map)+1}"
            criteria_map[key] = {"name": m.group(1).strip(), "max": int(m.group(2))}
    if not criteria_map:
        criteria_map = {"cat_1": {"name": "ความเหมาะสมโดยรวม", "max": 100}}
    return criteria_map


async def _ocr_file(filename: str) -> str:
    """Run OCR on a resume file and return extracted text, with caching."""
    path = os.path.join(RESUME_BASE, filename)
    if not os.path.exists(path):
        return ""

    # ─── Caching Logic ───
    safe_filename = "".join(c for c in filename if c.isalnum() or c in (".", "_", "-")).rstrip()
    cache_path = os.path.join(CACHE_DIR, f"ocr_{safe_filename}.txt")

    if os.path.exists(cache_path):
        src_mtime = os.path.getmtime(path)
        cache_mtime = os.path.getmtime(cache_path)
        if cache_mtime >= src_mtime:
            try:
                with open(cache_path, "r", encoding="utf-8") as f:
                    cached_text = f.read().strip()
                if cached_text:
                    logger.info(f"[OCR CACHE] Loaded cached OCR text for {filename} ({len(cached_text)} chars)")
                    return cached_text
            except Exception as e:
                logger.warning(f"Failed to read OCR cache for {filename}: {e}")

    with open(path, "rb") as f:
        content = f.read()

    ext = filename.lower()
    text = ""
    if ext.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            import io as _io
            reader = PdfReader(_io.BytesIO(content))
            text = "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception:
            text = ""

    elif ext.endswith(".docx"):
        try:
            from docx import Document
            import io as _io
            doc = Document(_io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            text = ""

    elif ext.endswith((".png", ".jpg", ".jpeg")):
        if "ocr_model" not in models:
            return f"OCR model not loaded for {filename}"
        try:
            import asyncio
            def sync_ocr():
                from PIL import Image
                import io as _io
                image = Image.open(_io.BytesIO(content)).convert("RGB")
                image = resize_if_needed(image)
                return _run_ocr_on_image(image, filename)
            
            text = await asyncio.to_thread(sync_ocr)
        except Exception as e:
            logger.exception(f"OCR failed for {filename}")
            text = ""
    else:
        try:
            text = content.decode("utf-8")
        except Exception:
            text = ""

    # Save to cache
    if text and len(text.strip()) > 10:
        try:
            with open(cache_path, "w", encoding="utf-8") as f:
                f.write(text)
            logger.info(f"[OCR CACHE] Saved OCR text for {filename} to cache")
        except Exception as e:
            logger.warning(f"Failed to write OCR cache for {filename}: {e}")

    return text


async def _score_resume(resume_text: str, jd_text: str, criteria_map: dict) -> dict:
    """Call Typhoon chat to score a resume and return structured result."""
    if "chat_model" not in models:
        return {"scores": {}, "detected_skills": [], "strengths": "", "summary": "Chat model not loaded"}

    cat_desc = ""
    catkeys = list(criteria_map.keys())
    for k in catkeys:
        cat_desc += f"- {k}: {criteria_map[k]['name']} (คะแนนเต็ม {criteria_map[k]['max']})\n"

    schema_keys = ", ".join(f'"{k}": <integer>' for k in catkeys)
    example_keys = ", ".join(f'"{k}": {int(criteria_map[k]["max"]*0.5)}' for k in catkeys)

    user_msg = f"""=== เกณฑ์การให้คะแนน (รวม 100 คะแนน) ===
{cat_desc}

=== เนื้อหา Resume ===
{resume_text}

=== Job Description ===
{jd_text}

=== กฎการให้คะแนน ===
1. ให้คะแนนตามเนื้อหาใน Resume จริงๆ เท่านั้น
2. ถ้าไม่มีหลักฐาน → 0 สำหรับหัวข้อนั้น
3. ตอบเป็น JSON เท่านั้น

JSON format: {{"scores":{{{schema_keys}}},"detected_skills":["skill1","skill2"],"strengths":"จุดเด่น","summary":"วิจารณ์สั้นๆ"}}"""

    system_msg = f"""You are a strict HR evaluator. Score resumes ONLY on evidence found in the text.
RULES: No hallucination. If skill not in resume → 0. Output ONLY valid JSON.
Example: {{"scores":{{{example_keys}}},"detected_skills":[],"strengths":"...","summary":"..."}}"""

    chat_req = ChatRequest(
        messages=[ChatMessage(role="user", content=user_msg)],
        system_prompt=system_msg,
        max_new_tokens=512,
        temperature=0.1,
    )

    # Collect streaming response
    from fastapi.responses import StreamingResponse as _SR
    resp = await chat_endpoint(chat_req)
    raw = b""
    async for chunk in resp.body_iterator:
        raw += chunk if isinstance(chunk, bytes) else chunk.encode()

    raw_str = raw.decode("utf-8", errors="ignore").strip()
    json_match = _re.search(r"\{[\s\S]*\}", raw_str)
    if not json_match:
        return {"scores": {k: 0 for k in catkeys}, "detected_skills": [], "strengths": "", "summary": "AI ไม่ส่ง JSON กลับมา"}

    try:
        return _json.loads(json_match.group(0))
    except Exception:
        return {"scores": {k: 0 for k in catkeys}, "detected_skills": [], "strengths": "", "summary": "JSON parse failed"}


@app.get("/api/analyze")
async def api_analyze(role: str = "fullstack"):
    """
    Llama-style: OCR all resumes then AI-score them for the given role.
    Returns list of results sorted by score descending.
    """
    jd_path  = os.path.join(JOBS_BASE, role, "jd.txt")
    cr_path  = os.path.join(JOBS_BASE, role, "criteria.txt")

    if not os.path.exists(jd_path) or not os.path.exists(cr_path):
        raise HTTPException(404, f"Role '{role}' not found")

    with open(jd_path,  "r", encoding="utf-8") as f:
        jd_text = f.read()
    with open(cr_path,  "r", encoding="utf-8") as f:
        criteria_text = f.read()

    criteria_map = _parse_criteria(criteria_text)

    if not os.path.exists(RESUME_BASE):
        return {"role": role, "job_title": jd_text.splitlines()[0], "results": []}

    files = [f for f in os.listdir(RESUME_BASE)
             if os.path.isfile(os.path.join(RESUME_BASE, f))
             and f.lower().endswith((".png", ".jpg", ".jpeg", ".pdf", ".docx", ".txt"))]

    if not files:
        return {"role": role, "job_title": jd_text.splitlines()[0], "results": [], "message": "No resumes found"}

    async def process_single(filename):
        logger.info(f"[API analyze] Processing: {filename}")
        resume_text = await _ocr_file(filename)

        if not resume_text or len(resume_text) < 20:
            return {
                "filename": filename,
                "score": 0,
                "breakdown": {criteria_map[k]["name"]: {"score": 0, "max": criteria_map[k]["max"]} for k in criteria_map},
                "detected_skills": [],
                "strengths": "",
                "summary": f"OCR ไม่สามารถอ่านข้อความได้ (ได้รับ {len(resume_text)} ตัวอักษร)",
                "error": "ocr_failed"
            }

        ai = await _score_resume(resume_text, jd_text, criteria_map)

        scores_raw = ai.get("scores", {})
        breakdown = {}
        total = 0.0
        for k, info in criteria_map.items():
            raw_val = scores_raw.get(k, 0)
            try:
                val = min(float(raw_val), info["max"])
            except (ValueError, TypeError):
                val = 0.0
            breakdown[info["name"]] = {"score": val, "max": info["max"]}
            total += val

        return {
            "filename": filename,
            "score": round(total, 1),
            "breakdown": breakdown,
            "detected_skills": ai.get("detected_skills", []),
            "strengths": ai.get("strengths", ""),
            "summary": ai.get("summary", ""),
        }

    import asyncio
    tasks = [process_single(filename) for filename in files]
    results = await asyncio.gather(*tasks)

    results.sort(key=lambda x: x.get("score", 0), reverse=True)

    return {
        "role": role,
        "job_title": jd_text.splitlines()[0] if jd_text.splitlines() else role,
        "results": results,
    }

